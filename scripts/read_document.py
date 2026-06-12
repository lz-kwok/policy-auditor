# -*- coding: utf-8 -*-
"""
read_document.py - Extracts text from .doc, .docx, .pdf, .xls, and .xlsx files to plain text for policy auditing.

Usage:
  python scripts/read_document.py --file <path_to_document> [--output <path_to_txt_file>]
"""

import os
import sys
import argparse

# Ensure stdout uses UTF-8 to prevent encoding issues with Chinese characters
try:
    sys.stdout.reconfigure(encoding='utf-8')
except AttributeError:
    pass

def read_docx(file_path):
    """Extracts text from a .docx file including paragraphs and tables."""
    try:
        import docx
    except ImportError:
        print("Error: 'python-docx' library is not installed. Please install it with: pip install python-docx")
        sys.exit(1)
    
    doc = docx.Document(file_path)
    full_text = []
    
    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
            
    # Extract tables
    for table in doc.tables:
        full_text.append("\n[表格数据]")
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            # Deduplicate adjacent identical cells due to cell merging
            dedup_row_text = []
            for item in row_text:
                if not dedup_row_text or item != dedup_row_text[-1]:
                    dedup_row_text.append(item)
            if dedup_row_text:
                full_text.append(" | ".join(dedup_row_text))
                
    return '\n'.join(full_text)

def read_doc_windows(file_path):
    """Extracts text from a .doc file on Windows using MS Word COM interface."""
    try:
        import win32com.client
    except ImportError:
        print("Error: 'pywin32' library is not installed. Please install it on Windows with: pip install pywin32")
        sys.exit(1)
        
    word = None
    doc = None
    try:
        # Get absolute path since Word COM requires it
        abs_path = os.path.abspath(file_path)
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(abs_path)
        text = doc.Content.Text
        # Clean up carriage returns
        text = text.replace('\r', '\n')
        return text
    except Exception as e:
        raise Exception(f"Failed to read .doc file via Word COM interface: {e}")
    finally:
        if doc:
            doc.Close(False)
        if word:
            word.Quit()

def read_pdf(file_path):
    """Extracts text from a .pdf file trying pypdf, pdfplumber, and fitz in order."""
    # 1. Try pypdf
    try:
        import pypdf
        reader = pypdf.PdfReader(file_path)
        text = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text.append(page_text)
        if text:
            return '\n'.join(text)
    except ImportError:
        pass
        
    # 2. Try pdfplumber
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            text = []
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
            if text:
                return '\n'.join(text)
    except ImportError:
        pass

    # 3. Try fitz (PyMuPDF)
    try:
        import fitz
        doc = fitz.open(file_path)
        text = []
        for page in doc:
            text.append(page.get_text())
        if text:
            return '\n'.join(text)
    except ImportError:
        pass
        
    print("Error: No PDF extraction library installed or text could not be extracted.")
    print("Please install one of the following libraries:")
    print("  pip install pypdf")
    print("  pip install pdfplumber")
    print("  pip install PyMuPDF")
    sys.exit(1)

def read_xlsx(file_path):
    """Extracts text from a .xlsx file sheet by sheet."""
    try:
        import openpyxl
    except ImportError:
        print("Error: 'openpyxl' library is not installed. Please install it with: pip install openpyxl")
        sys.exit(1)
        
    wb = openpyxl.load_workbook(file_path, data_only=True)
    text = []
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        text.append(f"\n[工作表: {sheet_name}]")
        for row in sheet.iter_rows(values_only=True):
            # Convert row cells to string, filtering out None values, and join with ' | '
            row_vals = [str(cell).strip() if cell is not None else "" for cell in row]
            if any(row_vals):  # only append non-empty rows
                text.append(" | ".join(row_vals))
    return '\n'.join(text)

def read_xls_xlrd(file_path):
    """Extracts text from a .xls file using xlrd."""
    try:
        import xlrd
    except ImportError:
        return None
        
    wb = xlrd.open_workbook(file_path)
    text = []
    for sheet in wb.sheets():
        text.append(f"\n[工作表: {sheet.name}]")
        for r in range(sheet.nrows):
            row_vals = [str(sheet.cell_value(r, c)).strip() for c in range(sheet.ncols)]
            if any(row_vals):
                text.append(" | ".join(row_vals))
    return '\n'.join(text)

def read_xls_windows(file_path):
    """Extracts text from a .xls file on Windows using Excel COM interface."""
    try:
        import win32com.client
    except ImportError:
        print("Error: 'pywin32' library is not installed. Please install it on Windows with: pip install pywin32")
        sys.exit(1)
        
    excel = None
    wb = None
    try:
        abs_path = os.path.abspath(file_path)
        excel = win32com.client.Dispatch("Excel.Application")
        excel.Visible = False
        wb = excel.Workbooks.Open(abs_path)
        text = []
        for sheet in wb.Sheets:
            text.append(f"\n[工作表: {sheet.Name}]")
            used_range = sheet.UsedRange
            rows_count = used_range.Rows.Count
            cols_count = used_range.Columns.Count
            for r in range(1, rows_count + 1):
                row_vals = []
                for c in range(1, cols_count + 1):
                    val = used_range.Cells(r, c).Value
                    row_vals.append(str(val).strip() if val is not None else "")
                if any(row_vals):
                    text.append(" | ".join(row_vals))
        return '\n'.join(text)
    except Exception as e:
        raise Exception(f"Failed to read .xls file via Excel COM interface: {e}")
    finally:
        if wb:
            wb.Close(False)
        if excel:
            excel.Quit()

def main():
    parser = argparse.ArgumentParser(description="Extract text from .doc, .docx, .pdf, .xls, and .xlsx files.")
    parser.add_argument("--file", required=True, help="Path to the document file")
    parser.add_argument("--output", help="Path to save the extracted plain text (optional)")
    args = parser.parse_args()
    
    file_path = args.file
    if not os.path.exists(file_path):
        print(f"Error: File not found: {file_path}")
        sys.exit(1)
        
    ext = os.path.splitext(file_path)[1].lower()
    
    try:
        if ext == ".docx":
            text = read_docx(file_path)
        elif ext == ".doc":
            if os.name == 'nt':
                text = read_doc_windows(file_path)
            else:
                print("Error: Reading .doc files is only supported on Windows via MS Word COM interface.")
                sys.exit(1)
        elif ext == ".pdf":
            text = read_pdf(file_path)
        elif ext == ".xlsx":
            text = read_xlsx(file_path)
        elif ext == ".xls":
            # Try xlrd first
            text = read_xls_xlrd(file_path)
            if text is None:
                if os.name == 'nt':
                    text = read_xls_windows(file_path)
                else:
                    print("Error: Reading .xls files requires 'xlrd' library on non-Windows platforms. Please run: pip install xlrd")
                    sys.exit(1)
        else:
            print(f"Error: Unsupported file format: {ext}. Only .doc, .docx, .pdf, .xls, and .xlsx are supported.")
            sys.exit(1)
            
        if args.output:
            out_dir = os.path.dirname(args.output)
            if out_dir and not os.path.exists(out_dir):
                os.makedirs(out_dir)
            with open(args.output, 'w', encoding='utf-8') as f:
                f.write(text)
            print(f"Text successfully extracted and saved to: {args.output}")
        else:
            print(text)
            
    except Exception as e:
        print(f"Error: {e}")
        sys.exit(1)

if __name__ == "__main__":
    main()
